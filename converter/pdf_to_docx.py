import tempfile, os
from io import BytesIO

def _convert_with_pdf2docx(pdf_bytes: bytes):
    try:
        from pdf2docx import Converter
    except ImportError:
        return None

    with tempfile.TemporaryDirectory() as td:
        pdf_path = os.path.join(td, "input.pdf")
        docx_path = os.path.join(td, "output.docx")
        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
        except Exception:
            return None
        if not os.path.exists(docx_path):
            return None
        with open(docx_path, "rb") as f:
            return f.read()

def _flow_text_and_images(pdf_bytes: bytes):
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        for i, page in enumerate(pdf):
            if i > 0:
                doc.add_page_break()
            title = doc.add_paragraph(f"Page {i+1}")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True

            text = page.get_text("text")
            for line in text.splitlines():
                doc.add_paragraph(line)

            for img in page.get_images(full=True):
                xref = img[0]
                pix = fitz.Pixmap(pdf, xref)
                if pix.alpha:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_bytes = pix.tobytes("png")
                doc.add_paragraph().add_run().add_picture(BytesIO(img_bytes), width=Inches(5))
                pix = None

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def convert_pdf_to_docx(input_bytes: bytes, strategy: str = "auto") -> bytes:
    strategy = (strategy or "auto").lower()
    if strategy in {"auto", "precise"}:
        out = _convert_with_pdf2docx(input_bytes)
        if out is not None:
            return out
        if strategy == "precise":
            raise RuntimeError("Precise conversion failed.")
    return _flow_text_and_images(input_bytes)
